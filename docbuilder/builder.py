from __future__ import annotations

import os
import re
from datetime import datetime

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
# pyrefly: ignore [missing-import]
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn

from agent.models import AgentState


def generate_docx(state: AgentState, output_path: str) -> None:
    """
    Generate a polished .docx document from agent task results.

    Parses task results as markdown and renders them into a styled
    Word document with cover page, headings, tables, and lists.
    """
    doc = Document()

    _set_default_style(doc)
    _add_cover_page(doc, state.request)
    _add_tasks_content(doc, state)
    _add_footer(doc)

    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)


def _set_default_style(doc: Document) -> None:
    """Configure default font and paragraph spacing."""
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)


def _add_cover_page(doc: Document, request: str) -> None:
    """Add a centered cover page with title and metadata."""
    for _ in range(6):
        doc.add_paragraph()

    title_text = _extract_title(request)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(title_text)
    run.bold = True
    run.font.size = Pt(26)
    run.font.color.rgb = RGBColor(0x1A, 0x23, 0x7E)
    run.font.name = "Calibri"

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Prepared by Autonomous AI Agent")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_p.add_run(datetime.now().strftime("%B %d, %Y"))
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    doc.add_page_break()


def _add_tasks_content(doc: Document, state: AgentState) -> None:
    """Render each task's result as a section in the document."""
    tasks = sorted(state.tasks, key=lambda t: t.id)

    for task in tasks:
        if task.status.value in ("failed", "skipped"):
            continue

        result_text = state.results.get(task.id, task.result)
        if not result_text:
            result_text = task.result if task.result else "[Content not available]"

        _render_markdown(doc, f"# {task.title}\n\n{result_text}")


def _render_markdown(doc: Document, markdown_text: str) -> None:
    """
    Parse a simplified markdown string and add matching docx elements.
    Supports: # h1, ## h2, ### h3, **bold**, - lists, |tables|.
    """
    lines = markdown_text.split("\n")
    in_table = False
    table_buffer = []

    i = 0
    while i < len(lines):
        line = lines[i]

        # Table detection
        if line.strip().startswith("|") and line.strip().endswith("|"):
            table_buffer.append(line.strip())
            in_table = True
            i += 1
            continue
        if in_table:
            _render_table(doc, table_buffer)
            table_buffer = []
            in_table = False
            continue

        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        # Heading 1
        if stripped.startswith("# ") and not stripped.startswith("## "):
            text = stripped[2:]
            h = doc.add_heading(text, level=1)
            _style_heading(h, 0)
            i += 1
            continue

        # Heading 2
        if stripped.startswith("## ") and not stripped.startswith("### "):
            text = stripped[3:]
            h = doc.add_heading(text, level=2)
            _style_heading(h, 1)
            i += 1
            continue

        # Heading 3
        if stripped.startswith("### "):
            text = stripped[4:]
            h = doc.add_heading(text, level=3)
            _style_heading(h, 2)
            i += 1
            continue

        # Bullet list
        if stripped.startswith("- ") or stripped.startswith("* "):
            text = stripped[2:]
            p = doc.add_paragraph(style="List Bullet")
            _add_run_with_bold(p, text)
            i += 1
            continue

        # Numbered list
        match = re.match(r"^(\d+)\.\s+(.*)", stripped)
        if match:
            text = match.group(2)
            p = doc.add_paragraph(style="List Number")
            _add_run_with_bold(p, text)
            i += 1
            continue

        # Paragraph
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(6)
        _add_run_with_bold(p, stripped)
        i += 1

    if in_table and table_buffer:
        _render_table(doc, table_buffer)


def _render_table(doc: Document, table_lines: list[str]) -> None:
    """Render markdown table lines as a styled docx table."""
    if len(table_lines) < 2:
        return

    header_cells = [c.strip() for c in table_lines[0].strip("|").split("|")]
    rows = []
    for line in table_lines[2:]:
        cells = [c.strip() for c in line.strip("|").split("|")]
        if cells:
            rows.append(cells)

    if not rows:
        return

    ncols = len(header_cells)
    table = doc.add_table(rows=1 + len(rows), cols=ncols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # Header row
    for j, cell_text in enumerate(header_cells):
        cell = table.rows[0].cells[j]
        cell.text = cell_text
        _shade_cell(cell, RGBColor(0x1A, 0x23, 0x7E))
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.size = Pt(10)

    # Data rows
    for i, row in enumerate(rows):
        for j, cell_text in enumerate(row):
            if j < ncols:
                cell = table.rows[i + 1].cells[j]
                cell.text = cell_text
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(10)
                if i % 2 == 1:
                    _shade_cell(cell, RGBColor(0xF5, 0xF5, 0xF5))

    doc.add_paragraph()


def _shade_cell(cell, color: RGBColor) -> None:
    """Apply background shading to a table cell."""
    shading = cell._element.get_or_add_tcPr()
    shading_elem = shading.makeelement(qn("w:shd"), {})
    shading_elem.set(qn("w:val"), "clear")
    shading_elem.set(qn("w:color"), "auto")
    shading_elem.set(qn("w:fill"), str(color))
    shading.append(shading_elem)


def _style_heading(heading, level: int) -> None:
    """Apply custom heading styling."""
    colors = [RGBColor(0x1A, 0x23, 0x7E), RGBColor(0x28, 0x35, 0x93), RGBColor(0x45, 0x45, 0x45)]
    sizes = [Pt(16), Pt(13), Pt(11)]
    if level < len(colors):
        for run in heading.runs:
            run.font.color.rgb = colors[level]
            run.font.size = sizes[level]
            run.font.name = "Calibri"


def _add_run_with_bold(paragraph, text: str) -> None:
    """Add text to a paragraph, handling **bold** markers."""
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)


def _add_footer(doc: Document) -> None:
    """Add a simple footer with agent branding."""
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("Generated by Autonomous AI Agent")
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)


def _extract_title(request: str) -> str:
    """Extract a short title from the request text."""
    cleaned = request.strip()
    if len(cleaned) > 80:
        cleaned = cleaned[:80] + "..."
    return cleaned
